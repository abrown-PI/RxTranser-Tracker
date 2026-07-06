from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x17, 0x2a, 0x4f)
    h.font.bold = True
    if level == 1:
        h.font.size = Pt(22)
    elif level == 2:
        h.font.size = Pt(16)
    else:
        h.font.size = Pt(13)

PI_BLUE = RGBColor(0x2a, 0x6e, 0xbb)
PI_DARK = RGBColor(0x17, 0x2a, 0x4f)
GRAY = RGBColor(0x47, 0x55, 0x69)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(text)
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)

def add_numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(text)
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)

def add_bold_body(bold_part, rest):
    p = doc.add_paragraph()
    r1 = p.add_run(bold_part)
    r1.bold = True
    p.add_run(rest)
    p.paragraph_format.space_after = Pt(6)

# ===== TITLE PAGE =====
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('PI Transfer Tracker', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('Quick Reference Guide')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in sub.runs:
    r.font.size = Pt(16)
    r.font.color.rgb = GRAY

doc.add_paragraph()
org = doc.add_paragraph('Pharmacy Innovations')
org.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in org.runs:
    r.font.size = Pt(14)
    r.font.color.rgb = PI_BLUE

doc.add_page_break()

# ===== OVERVIEW =====
doc.add_heading('Overview', level=2)
add_body('The PI Transfer Tracker is Pharmacy Innovations\u2019 tool for managing inter-pharmacy prescription transfers. It tracks prescriptions from the originating pharmacy location through sterile compounding at Erie and back to the origin or directly to the patient.')

# ===== GETTING STARTED =====
doc.add_heading('Getting Started', level=2)

doc.add_heading('Logging In', level=3)
add_bullet('In the current demo version, use the user switcher dropdown in the top-right of the navbar to switch between users.')
add_bullet('In production, authentication will be handled via Azure AD / MSAL.')

doc.add_heading('User Roles', level=3)
add_table(
    ['Role', 'Access'],
    [
        ['Admin', 'Full access to all locations, all tabs including Admin. Can manage users, SLA rules, and formula library.'],
        ['Manager', 'Full read access across locations plus Reports.'],
        ['Tech', 'Sees only transfers and shipments involving their own location (as origin or fill).'],
    ]
)

# ===== NAVIGATION TABS =====
doc.add_heading('Navigation Tabs', level=2)

# -- Transfers --
doc.add_heading('1. Transfers', level=3)
add_body('The main view showing all transfer requests visible to your role/location.')

add_bold_body('Filters available:', '')
add_bullet('Search by patient name, Rx number, or tracking number')
add_bullet('Filter by status, origin location, or fill location')

add_bold_body('Transfer Statuses:', '')
add_table(
    ['Status', 'Meaning'],
    [
        ['New', 'Just submitted, not yet picked up by the fill location'],
        ['Tx Verified', 'Transfer paperwork has been verified against the hard copy'],
        ['In Progress', 'Actively being compounded at the fill location'],
        ['Ready to Ship', 'Compounding complete, awaiting shipment'],
        ['Shipped', 'Package is in transit'],
        ['Delivered', 'Patient or pharmacy has received the order'],
        ['Canceled', 'Transfer was canceled'],
    ]
)

add_bold_body('Clicking a transfer', ' opens its detail modal where you can:')
add_bullet('Update the status')
add_bullet('View/edit shipping and tracking info')
add_bullet('Review attached documents (PK Sheet, Formula, Hard Copy)')
add_bullet('Verify prescriptions against hard copies (checkbox per item)')
add_bullet('Post comments (scoped to the whole transfer or a specific Rx)')

# -- New Transfer --
doc.add_heading('2. + New Transfer', level=3)
add_body('Submit a new transfer request. Required fields are marked with a red asterisk.')

add_bold_body('Required:', '')
add_bullet('Patient Name')
add_bullet('At least one prescription with a drug name')
add_bullet('PK Transfer Sheet (upload PDF/image)')
add_bullet('Formula Worksheet (select from location library OR upload new)')

add_bold_body('Key sections:', '')
add_bullet('Name, DOB, allergies', bold_prefix='Patient \u2014 ')
add_bullet('Origin location (your pharmacy) and fill location (typically Erie for sterile compounds)', bold_prefix='Routing \u2014 ')
add_bullet('Ship to Pharmacy, Patient, or Delivery; address fields; cold item checkbox', bold_prefix='Shipping \u2014 ')
add_bullet('HealNow, Credit Card, or Pay In Store; Sema Dr Authorization status', bold_prefix='Billing \u2014 ')
add_bullet('PK Transfer Sheet, Formula (from library or upload), optional Hard Copy scan', bold_prefix='Documents \u2014 ')
add_bullet('Drug, quantity, Rx number, doctor info, receipt attached, notes. Add multiple Rx lines with "+ Add another Rx"', bold_prefix='Items/Prescriptions \u2014 ')

add_bold_body('Formula Library: ', 'When creating a transfer, you can pick a saved formula from your location\u2019s library instead of uploading. If you upload a new formula, check "Save to my location\u2019s library for reuse" to add it for future use.')

# -- Pharmacist Review --
doc.add_heading('3. Pharmacist Review', level=3)
add_body('Shows transfers that have unverified items \u2014 prescriptions where the PK transfer form has not yet been confirmed against the hard copy. A badge on the tab shows the count.')
add_bullet('Admins see all unverified items')
add_bullet('Fill-location techs see items they need to verify')
add_bullet('Check the "PK form matches hard copy" box on each item once verified')

# -- Shipments --
doc.add_heading('4. Shipments', level=3)
add_body('Manage physical shipments between locations.')

add_bold_body('Features:', '')
add_bullet('Focus the scan field and scan a tracking barcode to jump directly to that shipment', bold_prefix='Barcode scanner input \u2014 ')
add_bullet('Bundle one or more transfers into a shipment with carrier/tracking info', bold_prefix='Create Shipment \u2014 ')
add_bullet('When a shipment arrives at your location, click "Receive" to confirm arrival, unpack and verify each item, and report any damage (with photo upload)', bold_prefix='Receive Shipment \u2014 ')
add_bullet('Filter by your location or search by tracking number')

add_bold_body('Shipment Statuses:', '')
add_table(
    ['Status', 'Meaning'],
    [
        ['Pending', 'Shipment created but not yet sent'],
        ['In Transit', 'Shipped and on its way'],
        ['Delivered', 'Carrier confirms delivery'],
        ['Received', 'Staff at the destination has physically received and unpacked'],
    ]
)

# -- SLA Alerts --
doc.add_heading('5. SLA Alerts', level=3)
add_body('Flags transfers that are approaching or have exceeded their time limits.')

add_bold_body('Default SLA Rules:', '')
add_table(
    ['Status', 'Time Limit'],
    [
        ['New', '24 hours'],
        ['Tx Verified', '48 hours'],
        ['In Progress', '72 hours'],
        ['Ready to Ship', '24 hours'],
    ]
)
add_bullet('Transfer has exceeded the allowed time in its current status', bold_prefix='Breach (red) \u2014 ')
add_bullet('Transfer is at 75%+ of the allowed time', bold_prefix='Near SLA (yellow) \u2014 ')
add_bullet('SLA rules can be adjusted by admins in the Admin tab')
add_bullet('In production, breaches will auto-notify assigned users via Teams')

# -- Reports --
doc.add_heading('6. Reports', level=3)
add_body('Dashboard with KPIs and analytics.')

add_bold_body('KPI Cards:', '')
add_bullet('Total Transfers (scoped to your visibility)')
add_bullet('In Flight (not yet delivered)')
add_bullet('Avg Cycle Time (Tx email to ship date)')
add_bullet('Avg Turnaround (Tx email to delivered)')
add_bullet('On-Time Ship Rate')
add_bullet('SLA Breaches')
add_bullet('Cold Chain Items')

add_bold_body('Report Tables:', '')
add_bullet('Volume by Status')
add_bullet('Avg Turnaround by Origin Location')
add_bullet('Volume by Origin Location')
add_bullet('Unverified Items list')

add_bold_body('Export: ', 'Click "Export CSV" to download all visible transfer data.')

# -- Admin --
doc.add_heading('7. Admin (Admin role only)', level=3)

add_bold_body('Users & Roles', '')
add_bullet('Add or remove users')
add_bullet('Change a user\u2019s location or role (admin / manager / tech)')

add_bold_body('SLA Rules', '')
add_bullet('Adjust the hour limits for each status threshold')
add_bullet('Click "Save SLA rules" to apply')

add_bold_body('Formula Library', '')
add_bullet('Add, edit, or remove reusable formula worksheets')
add_bullet('Each formula is scoped to a specific location')
add_bullet('Formulas can include a SharePoint link for the source document')

add_bold_body('Teams Webhooks', '')
add_bullet('Configure Microsoft Teams webhook URLs per location for automated notifications (planned for production)')

# ===== GLOBAL SEARCH =====
doc.add_heading('Global Search', level=2)
add_body('The search bar in the top navbar searches across all visible transfers and shipments by:')
add_bullet('Patient name')
add_bullet('Rx number')
add_bullet('Tracking number')
add_bullet('Doctor name')
add_bullet('Location')
add_bullet('Drug name')
add_body('Click a result to jump directly to that transfer or shipment.')

# ===== KEY WORKFLOWS =====
doc.add_heading('Key Workflows', level=2)

doc.add_heading('Submitting a Transfer (Origin Pharmacy Tech)', level=3)
add_numbered('Go to + New Transfer')
add_numbered('Fill in patient info, routing, shipping, and billing')
add_numbered('Upload the PK Transfer Sheet')
add_numbered('Select or upload a Formula Worksheet')
add_numbered('Add prescription line items with doctor details')
add_numbered('Click Submit transfer')

doc.add_heading('Processing a Transfer (Fill Location / Erie)', level=3)
add_numbered('Check the Transfers tab for new incoming transfers')
add_numbered('Open the transfer and review documents and items')
add_numbered('Verify each item against the hard copy (check the box)')
add_numbered('Update status: New \u2192 Tx Verified \u2192 In Progress \u2192 Ready to Ship')
add_numbered('Create a shipment when ready')

doc.add_heading('Receiving a Shipment (Destination Pharmacy Tech)', level=3)
add_numbered('Go to Shipments tab')
add_numbered('Scan the tracking barcode or find the shipment in the list')
add_numbered('Click Receive')
add_numbered('Verify each item and report any damage')
add_numbered('Shipment moves to "Received" status')

# ===== TIPS =====
doc.add_heading('Tips', level=2)
add_bullet('are flagged with a blue "Cold" badge \u2014 these require special packaging and handling', bold_prefix='Cold items ')
add_bullet('in the Transfers table indicate SLA status (red = breach, yellow = near SLA)', bold_prefix='Color-coded rows ')
add_bullet('can be scoped to a specific Rx item or the whole transfer \u2014 use the dropdown before posting', bold_prefix='Comments ')
add_bullet('checkbox on comments will send a Teams notification in production', bold_prefix='The "Notify via Teams" ')
add_bullet('show a yellow warning pill in the Transfers table and appear in Pharmacist Review', bold_prefix='Unverified items ')

# ===== LOCATIONS =====
doc.add_heading('Locations', level=2)
add_body('Erie  \u2022  Lancaster  \u2022  Greenville  \u2022  Spring  \u2022  Tucson  \u2022  Flower Mound  \u2022  Denton  \u2022  Corinth  \u2022  Jamestown  \u2022  Virginia Beach  \u2022  Seminole')

# -- Footer --
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PI Transfer Tracker \u2014 Pharmacy Innovations')
r.font.size = Pt(10)
r.font.color.rgb = GRAY
r.italic = True

# Save
out_path = r'\\wsl.localhost\Ubuntu\home\abrown\transfer-tracker\PI-Transfer-Tracker-Reference-Guide.docx'
doc.save(out_path)
print(f'Saved to {out_path}')
